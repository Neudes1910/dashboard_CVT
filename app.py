import streamlit as st
import pandas as pd
from docx import Document
import io

st.set_page_config(page_title="Analisador Automático de Relatórios - CVT", layout="wide")
st.title("Analisador Automático de Relatórios - CVT")

uploaded_files = st.file_uploader(
    "Envie os relatórios CVT (.docx ou .docm)",
    type=["docx","docm"],
    accept_multiple_files=True
)

def documento_tem_hidrometer(doc):
    for p in doc.paragraphs:
        if "hidrometer connect" in p.text.lower():
            return True

    for tabela in doc.tables:
        for linha in tabela.rows:
            for cel in linha.cells:
                if "hidrometer connect" in cel.text.lower():
                    return True
    return False


def encontrar_tabela_ocorrencias(doc):

    for tabela in doc.tables:

        cabecalho = [c.text.strip().upper() for c in tabela.rows[0].cells]

        if "NATUREZA" in cabecalho and "OCORRÊNCIA" in cabecalho:

            idx_natureza = cabecalho.index("NATUREZA")

            registros = []

            for linha in tabela.rows[1:]:

                natureza = linha.cells[idx_natureza].text.strip()

                if natureza and natureza.lower() not in [
                    "escolha um item",
                    "escolher um item."
                ]:
                    registros.append(natureza)

            return registros

    return []


def encontrar_tabela_indisponibilidade(doc):

    registros = []

    for tabela in doc.tables:

        cabecalho = [c.text.strip().upper() for c in tabela.rows[0].cells]

        if "POR QUANTO TEMPO?" in cabecalho and "QUAL EQUIPAMENTO?" in cabecalho:

            idx_horas = cabecalho.index("POR QUANTO TEMPO?")
            idx_equip = cabecalho.index("QUAL EQUIPAMENTO?")
            idx_nat = cabecalho.index("NATUREZA")

            for linha in tabela.rows[1:]:

                horas_txt = linha.cells[idx_horas].text.strip()
                equipamento = linha.cells[idx_equip].text.strip()
                natureza = linha.cells[idx_nat].text.strip()

                try:
                    horas = float(horas_txt.replace(",","."))
                except:
                    horas = 0

                if natureza.lower() in [
                    "escolha um item",
                    "escolher um item."
                ]:
                    continue

                registros.append({
                    "equipamento": equipamento,
                    "natureza": natureza,
                    "horas": horas
                })

    return registros


if uploaded_files:

    ocorrencias = []
    indisponibilidades = []

    for arquivo in uploaded_files:

        try:
            doc = Document(io.BytesIO(arquivo.read()))
        except:
            continue

        if not documento_tem_hidrometer(doc):
            continue

        ocorrencias += encontrar_tabela_ocorrencias(doc)
        indisponibilidades += encontrar_tabela_indisponibilidade(doc)

    if not ocorrencias and not indisponibilidades:
        st.warning("Nenhum registro encontrado nos arquivos.")
        st.stop()

    col1, col2 = st.columns(2)

    # -------- OCORRÊNCIAS --------

    if ocorrencias:

        df_oc = pd.DataFrame({"Natureza": ocorrencias})

        contagem_nat = df_oc.value_counts().reset_index()
        contagem_nat.columns = ["Natureza","Total"]

        with col1:
            st.subheader("Total de Ocorrências por Natureza")
            st.bar_chart(contagem_nat.set_index("Natureza"))

    # -------- INDISPONIBILIDADE --------

    if indisponibilidades:

        df_ind = pd.DataFrame(indisponibilidades)

        total_horas = df_ind["horas"].sum()

        st.subheader("Total de Horas Indisponíveis")
        st.metric("Horas totais", round(total_horas,2))

        col3, col4 = st.columns(2)

        # horas por natureza

        horas_nat = df_ind.groupby("natureza")["horas"].sum().reset_index()

        with col3:
            st.subheader("Horas Indisponíveis por Natureza")
            st.bar_chart(horas_nat.set_index("natureza"))

        # horas por equipamento (CORRETO)

        horas_eq = df_ind.groupby("equipamento")["horas"].sum().reset_index()

        with col4:
            st.subheader("Horas Indisponíveis por Equipamento")
            st.bar_chart(horas_eq.set_index("equipamento"))

        st.plotly_chart(fig3, use_container_width=True)

else:
    st.info("Aguardando envio dos relatórios.")

